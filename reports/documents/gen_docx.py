#!/usr/bin/env python3
"""
Convert smart_search_technical_report.md → smart_search_technical_report.docx

Features:
  - Embedded PNG images with captions at the positions marked in the markdown
  - Word TOC field (right-click → Update Field for live page numbers + hyperlinks)
  - Proper Word Heading styles (Heading 1/2/3) so navigation pane and TOC work
  - Page break before every numbered top-level section (## 1., ## 2., ...)
  - Numbered and bulleted lists
  - Code blocks with grey background, monospace font
  - Inline bold / italic / code
  - Tables with blue header row
"""

import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH   = "smart_search_technical_report.md"
DOCX_PATH = "smart_search_technical_report.docx"

# ── Colours ────────────────────────────────────────────────────────────────────
H1_COLOR   = RGBColor(0x1F, 0x49, 0x7D)   # dark navy  — document title
H2_COLOR   = RGBColor(0x1D, 0x4E, 0xD8)   # blue       — ## headings  (Heading 1)
H3_COLOR   = RGBColor(0x15, 0x80, 0x3D)   # green      — ### headings (Heading 2)
H4_COLOR   = RGBColor(0x6D, 0x28, 0xD9)   # purple     — #### headings (Heading 3)
CAP_COLOR  = RGBColor(0x47, 0x55, 0x69)   # grey       — image captions
CODE_BG    = "F2F2F2"                       # light grey — code blocks / inline code


# ── XML / formatting helpers ───────────────────────────────────────────────────

def _shd(element, fill_hex):
    """Attach a w:shd element with the given fill colour to any XML node."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    element.append(shd)


def set_cell_bg(cell, fill_hex):
    _shd(cell._tc.get_or_add_tcPr(), fill_hex)


def set_run_bg(run, fill_hex):
    _shd(run._r.get_or_add_rPr(), fill_hex)


def add_toc_field(doc):
    """
    Insert a Word TOC field.  When the .docx is opened in Word, right-click the
    placeholder text and choose 'Update Field' to populate headings + page numbers
    with live hyperlinks.
    """
    # Heading for the TOC section
    h = doc.add_heading("Table of Contents", level=1)
    if h.runs:
        h.runs[0].font.color.rgb = H2_COLOR
        h.runs[0].font.size = Pt(14)

    para = doc.add_paragraph()
    run  = para.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'   # levels 1-3, with hyperlinks

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")

    # Placeholder text shown before the field is updated
    placeholder_r = OxmlElement("w:r")
    placeholder_t = OxmlElement("w:t")
    placeholder_t.text = (
        "[Right-click this line in Word and choose "
        "‘Update Field’ to build the Table of Contents]"
    )
    placeholder_r.append(placeholder_t)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(sep)
    run._r.append(placeholder_r)
    run._r.append(end)

    run.italic = True
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    run.font.size = Pt(10)


def add_code_block(doc, code_text):
    """Grey-background monospaced paragraph for fenced code blocks."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Inches(0.3)
    para.paragraph_format.right_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(6)
    _shd(para._p.get_or_add_pPr(), CODE_BG)
    run = para.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)


def add_image(doc, base_dir, filename, caption):
    """Centre the image and add an italic caption below it."""
    full = os.path.join(base_dir, filename)
    if not os.path.exists(full):
        doc.add_paragraph(f"[Image not found: {filename}]")
        return

    # Vertical spacing before
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)

    # Image — centred, max width 6.2 inches
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.add_run().add_picture(full, width=Inches(6.2))

    # Caption — centred, smaller, grey, italic
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_para.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = CAP_COLOR

    # Vertical spacing after
    sp2 = doc.add_paragraph()
    sp2.paragraph_format.space_before = Pt(4)
    sp2.paragraph_format.space_after  = Pt(8)


def add_table(doc, table_lines):
    """Render a markdown pipe-table as a Word table with a blue header row."""
    # Drop separator rows (---|---|---)
    rows = [l for l in table_lines
            if not re.match(r'^\s*\|[-| :]+\|\s*$', l)]
    if not rows:
        return

    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)

    cols = max(len(r) for r in parsed)
    tbl  = doc.add_table(rows=len(parsed), cols=cols)
    tbl.style = "Table Grid"

    for ri, row in enumerate(parsed):
        for ci in range(cols):
            text = row[ci] if ci < len(row) else ""
            cell = tbl.cell(ri, ci)
            cell.text = ""
            para = cell.paragraphs[0]
            apply_inline(para, text, base_size=9.5)
            if ri == 0:                              # header row
                for run in para.runs:
                    run.bold = True
                set_cell_bg(cell, "D6E4F0")          # pale blue header


def apply_inline(para, text, base_size=11):
    """
    Parse **bold**, *italic*, `code` markers and append styled runs to *para*.
    Everything else is plain text at *base_size*.
    """
    pattern = r"(\*\*[^*]+?\*\*|\*[^*]+?\*|`[^`]+?`)"
    for chunk in re.split(pattern, text):
        if not chunk:
            continue
        if re.fullmatch(r"\*\*(.+)\*\*", chunk):
            r = para.add_run(chunk[2:-2])
            r.bold = True
            r.font.size = Pt(base_size)
        elif re.fullmatch(r"\*(.+)\*", chunk):
            r = para.add_run(chunk[1:-1])
            r.italic = True
            r.font.size = Pt(base_size)
        elif re.fullmatch(r"`([^`]+)`", chunk):
            r = para.add_run(chunk[1:-1])
            r.font.name = "Courier New"
            r.font.size = Pt(base_size - 1)
            set_run_bg(r, CODE_BG)
        else:
            r = para.add_run(chunk)
            r.font.size = Pt(base_size)


# ── Main conversion ────────────────────────────────────────────────────────────

def convert(md_path, docx_path):
    base_dir = os.path.dirname(os.path.abspath(md_path))

    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()

    # Default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # Margins: 1.25 in sides, 1 in top/bottom
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    in_code      = False
    code_lines   = []
    skip_toc     = False   # True while consuming manual TOC list items in MD

    i = 0
    while i < len(lines):
        raw  = lines[i].rstrip("\n")
        line = raw.strip()

        # ── Fenced code block ─────────────────────────────────────────────────
        if line.startswith("```"):
            if not in_code:
                in_code    = True
                code_lines = []
            else:
                add_code_block(doc, "\n".join(code_lines))
                in_code = False
            i += 1
            continue

        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        # ── Image  ![Caption](file.png) ───────────────────────────────────────
        m = re.match(r"^!\[([^\]]*)\]\(([^)]+\.png)\)$", line)
        if m:
            add_image(doc, base_dir, m.group(2), m.group(1))
            i += 1
            continue

        # ── Table ─────────────────────────────────────────────────────────────
        if line.startswith("|"):
            tbl_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i].strip())
                i += 1
            add_table(doc, tbl_lines)
            doc.add_paragraph()
            continue

        # ── Horizontal rule ───────────────────────────────────────────────────
        if re.match(r"^---+$", line):
            i += 1
            continue

        # ── Headings ──────────────────────────────────────────────────────────
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text  = m.group(2).strip()
            skip_toc = False

            # ── ## Table of Contents → replace with Word TOC field ───────────
            if text == "Table of Contents":
                add_toc_field(doc)
                skip_toc = True
                i += 1
                continue

            # Page break before numbered top-level sections (## 1. … ## 13.)
            if level == 2 and re.match(r"^\d+\.", text):
                doc.add_page_break()

            if level == 1:                              # # Title
                para = doc.add_heading(text, level=0)
                if para.runs:
                    para.runs[0].font.color.rgb = H1_COLOR
                    para.runs[0].font.size = Pt(20)
            elif level == 2:                            # ## → Heading 1
                para = doc.add_heading(text, level=1)
                if para.runs:
                    para.runs[0].font.color.rgb = H2_COLOR
                    para.runs[0].font.size = Pt(14)
            elif level == 3:                            # ### → Heading 2
                para = doc.add_heading(text, level=2)
                if para.runs:
                    para.runs[0].font.color.rgb = H3_COLOR
                    para.runs[0].font.size = Pt(12)
            else:                                       # #### → Heading 3
                para = doc.add_heading(text, level=3)
                if para.runs:
                    para.runs[0].font.color.rgb = H4_COLOR
                    para.runs[0].font.size = Pt(11)

            i += 1
            continue

        # ── Skip manual TOC list lines after "## Table of Contents" ──────────
        if skip_toc:
            if re.match(r"^\d+\.", line) or line == "":
                i += 1
                continue
            else:
                skip_toc = False   # real content starts → stop skipping

        # ── Unordered list ─────────────────────────────────────────────────────
        if re.match(r"^[-*]\s+", line):
            para = doc.add_paragraph(style="List Bullet")
            apply_inline(para, re.sub(r"^[-*]\s+", "", line))
            i += 1
            continue

        # ── Ordered list ───────────────────────────────────────────────────────
        if re.match(r"^\d+\.\s+", line):
            para = doc.add_paragraph(style="List Number")
            apply_inline(para, re.sub(r"^\d+\.\s+", "", line))
            i += 1
            continue

        # ── Blank line ─────────────────────────────────────────────────────────
        if line == "":
            i += 1
            continue

        # ── Regular paragraph ──────────────────────────────────────────────────
        para = doc.add_paragraph()
        apply_inline(para, line)
        i += 1

    doc.save(docx_path)
    print(f"Saved → {docx_path}")


if __name__ == "__main__":
    base = os.path.dirname(os.path.abspath(__file__))
    convert(
        os.path.join(base, MD_PATH),
        os.path.join(base, DOCX_PATH),
    )
