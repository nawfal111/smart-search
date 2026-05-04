"""
generate_report.py
Generates report.docx — the full Smart Search technical research report.
Run: python generate_report.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Styles helpers ─────────────────────────────────────────────────────────────
def set_font(run, size=11, bold=False, italic=False, color=None, mono=False):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if mono:
        run.font.name = "Courier New"
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    p.runs[0].font.size = Pt(20)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x25, 0x63, 0xeb)
    p.runs[0].font.size = Pt(15)
    return p

def heading3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x37, 0x37, 0x55)
    p.runs[0].font.size = Pt(13)
    return p

def body(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r, size=11)
    return p

def numbered(text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_font(r, size=11)
    return p

def code_block(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent  = Inches(0.3)
    pf.space_before = Pt(6)
    pf.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x20, 0x20, 0x40)
    # light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F3F4F6")
    pPr.append(shd)
    return p

def diagram(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent  = Inches(0.2)
    pf.space_before = Pt(8)
    pf.space_after  = Pt(8)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x1e, 0x1e, 0x3e)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F4FF")
    pPr.append(shd)
    return p

def callout(label, text, color="0050AA"):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent  = Inches(0.3)
    pf.space_before = Pt(4)
    pf.space_after  = Pt(4)
    r1 = p.add_run(label.upper() + "  ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor.from_string(color)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for run in cells[i].paragraphs[0].runs:
                run.font.size = Pt(10)
    doc.add_paragraph()

def page_break():
    doc.add_page_break()

def divider():
    p = doc.add_paragraph("─" * 90)
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)

def challenge_header(num, title):
    p = doc.add_paragraph()
    r1 = p.add_run(f"[{num}]  ")
    r1.bold = True
    r1.font.size = Pt(13)
    r1.font.color.rgb = RGBColor(0x25, 0x63, 0xeb)
    r2 = p.add_run(title)
    r2.bold = True
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    divider()

# ══════════════════════════════════════════════════════════════════ COVER ══
heading1("Smart Search")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A Semantic Code Search VS Code Extension\nBuilt on Voyage AI · Pinecone · OpenAI GPT")
r.italic = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
add_table(
    ["Field", "Value"],
    [
        ["Author", "Nawfal Jalloul"],
        ["Purpose", "Personal productivity tool for navigating large codebases"],
        ["Extension Stack", "TypeScript · VS Code Extension API · Node.js"],
        ["Backend Stack", "Python · HTTP Server · ThreadPoolExecutor"],
        ["AI Services", "Voyage AI (voyage-code-2) · OpenAI (gpt-4o-mini) · Pinecone"],
        ["Languages Supported", "Python · TypeScript · JavaScript · PHP · Java · Go · Rust · C# · Ruby · Swift · Kotlin · C/C++"],
        ["Date", "May 2026"],
    ]
)
page_break()

# ══════════════════════════════════════════════════════════════ ABSTRACT ══
heading1("Abstract")
body(
    "Smart Search is a Visual Studio Code extension that adds AI-powered semantic code search "
    "to any project. It lets a developer type a natural language question — \"where do we validate "
    "the payment subscription?\" — and receive the exact function and specific line number that "
    "answers it, even when no keyword from the query appears in the code. The system indexes source "
    "code by splitting it into function-level chunks, generating plain English summaries via GPT, "
    "embedding those summaries alongside the code with Voyage AI's code-specialized model, and "
    "storing the resulting vectors in Pinecone. Search queries are embedded at query time and matched "
    "via cosine similarity. A second LLM call pinpoints the most relevant line within each matched "
    "function. This paper documents the complete system design, all engineering challenges encountered "
    "during development, and the solutions implemented. It serves as a technical record of this "
    "personal productivity project."
)
page_break()

# ══════════════════════════════════════════════════════════ 01 MOTIVATION ══
heading1("01  Motivation and Personal Context")
body(
    "The problem that prompted this project is one every developer working on a medium-to-large "
    "codebase encounters daily: you know a piece of logic exists somewhere, but you cannot remember "
    "what it is called."
)
body(
    "In a typical web application backend — a Laravel PHP project with 400 files, or a Node.js API "
    "with 200 TypeScript modules — the built-in VS Code search (Ctrl+Shift+F) and grep are the "
    "primary navigation tools. Both require you to know a keyword. If you know the function is named "
    "checkSubscription, grep will find it instantly. But if you only know what it does — \"checks "
    "whether a user's subscription is still active\" — you are stuck. You must either browse files "
    "manually, ask a colleague, or guess at keywords and hope."
)
body("This was a recurring personal pain point. The specific triggers were:")
bullet("Returning to a codebase after weeks away and needing to find logic written months earlier under an unclear name.")
bullet("Inheriting or reviewing a project written by someone else with different naming conventions.")
bullet("Looking for where a specific business rule is enforced in a large multi-module system.")

heading2("Core Goal")
callout("GOAL", "Type a natural language question into VS Code. Get back the exact file, function, and line number — even when no word from your question appears in the code.")

heading2("Design Constraints")
add_table(
    ["Constraint", "Implication"],
    [
        ["Must work inside VS Code", "VS Code Extension API, webview panel for UI"],
        ["Must not slow editor startup", "Indexing runs in background, non-blocking"],
        ["Must not re-embed unchanged code", "Two-level MD5 hash comparison before any API call"],
        ["Must support multiple languages", "11 language-specific chunkers in TypeScript"],
        ["Must isolate each user's vectors", "Pinecone namespace = projectId :: MD5(git email)"],
        ["Must not require Python on every machine", "Configurable backend URL (can point to remote server)"],
        ["Must keep source code off third-party servers", "Only vectors and short metadata go to Pinecone"],
    ]
)
page_break()

# ══════════════════════════════════════════════════════ 02 ARCHITECTURE ══
heading1("02  System Architecture Overview")
body(
    "The system is divided into two major layers that communicate over HTTP, plus three external "
    "cloud services."
)

diagram(
"""╔══════════════════════════════════════════════════════════════════╗
║                   DEVELOPER'S MACHINE                            ║
║                                                                  ║
║  ┌──────────────────────────────────────────────────────────┐   ║
║  │             VS CODE EXTENSION (TypeScript)               │   ║
║  │                                                          │   ║
║  │  Webview   WorkspaceIndexer   FileWatcher   Replace      │   ║
║  │  Panel     (chunker/hash)     (save/del/    Handler      │   ║
║  │  (HTML/JS) localIndex         rename)       (WsEdit)     │   ║
║  │            projectId/userId                              │   ║
║  └──────┬─────────────┬──────────────────────────────────-──┘   ║
║         │ postMessage  │ HTTP POST /index                        ║
║         │ search req   │                                         ║
║         ▼              ▼                                         ║
║  ┌──────────────────────────────────────────────────────────┐   ║
║  │              PYTHON BACKEND (server.py)                  │   ║
║  │                                                          │   ║
║  │  GET  /health  → { ok: true }                            │   ║
║  │  POST /search  → search.py (normal)                      │   ║
║  │               → embedder + pinecone + line_locator (AI)  │   ║
║  │  POST /index   → summarizer → embedder → pinecone        │   ║
║  │  POST /wipe    → pinecone_client.wipe_namespace()         │   ║
║  └──────────────────────────────────────────────────────────┘   ║
║         │                │               │                       ║
╚═════════╪════════════════╪═══════════════╪═══════════════════════╝
          │                │               │
    Voyage AI         OpenAI GPT       Pinecone
    voyage-code-2     gpt-4o-mini      Vector DB
    embeddings        summarizer       (cloud)
                      line_locator"""
)

heading2("Separation of Concerns")
add_table(
    ["Layer", "Responsibilities", "Why this layer?"],
    [
        ["TypeScript Extension", "File watching, chunking, hashing, UI, VS Code API calls, replace", "Must run inside VS Code; has native VS Code API access"],
        ["Python Backend", "LLM calls (GPT), embeddings (Voyage AI), Pinecone, regex search", "Mature AI/ML SDKs; I/O-bound parallel calls with ThreadPoolExecutor"],
    ]
)
body(
    "The backend is stateless — it holds no session data and no index. All state lives either in "
    "Pinecone (vectors) or in the local .smart-search/index.json file on the developer's machine (hashes)."
)
page_break()

# ══════════════════════════════════════════════════════ 03 COMPONENTS ══
heading1("03  Component Deep Dive")

heading2("3.1 — VS Code Extension Layer")
body("The extension source is organized as follows:")
diagram(
"""src/
├── extension.ts          ← Entry point. Wires up all 8 lifecycle events.
├── config.ts             ← getBackendUrl() — reads smartSearch.backendUrl setting
├── handlers/
│   ├── searchHandler.ts  ← POSTs to /search, returns results to webview
│   └── replaceHandler.ts ← VS Code WorkspaceEdit for single and bulk replace
├── indexer/
│   ├── workspaceIndexer.ts  ← Two-level hash logic; drives all indexing flows
│   ├── chunker.ts           ← Splits code into function chunks (11 languages)
│   ├── localIndex.ts        ← Reads/writes .smart-search/index.json
│   ├── projectId.ts         ← Generates/reads UUID in .smart-search/project-id
│   └── userId.ts            ← MD5(git global email) for Pinecone namespace
└── utils/
    └── webviewManager.ts    ← Inlines CSS+JS into HTML for the webview panel"""
)

heading3("extension.ts — Entry Point")
body("The activate() function registers 8 lifecycle behaviours, all non-blocking:")
numbered("Status bar item — shows indexing progress at the bottom of VS Code")
numbered("Backend health check — GETs /health; shows a warning if Python server is not running")
numbered("Workspace indexing — walks all files in the background on every VS Code open")
numbered("Save listener — re-indexes a single file every time it is saved")
numbered("Delete listener — removes deleted files' vectors from Pinecone immediately")
numbered("Rename listener — removes old path vectors, re-indexes file at new path")
numbered("Re-index command — wipes Pinecone + rebuilds from scratch")
numbered("Search panel command — opens the webview UI panel")

heading3("chunker.ts — The Code Splitter")
body(
    "The chunker takes a file's raw text and returns an array of Chunk objects. Each chunk represents "
    "one function, method, or class. It runs entirely in TypeScript — no network call. "
    "Chunking a 500-line file takes under 1ms."
)
body("The Chunk interface:")
code_block(
"""{ id:         "src/auth.php::verifyToken",  // unique: relativePath::name
  name:       "verifyToken",
  type:       "function",                   // "function"|"class"|"method"|"file"
  content:    "function verifyToken($tok) { ... }",
  start_line: 12,
  end_line:   38,
  file:       "src/auth.php",              // relative path (portable)
  language:   "php" }"""
)

heading3("workspaceIndexer.ts — The Brain")
body(
    "Orchestrates everything. Implements two levels of hash comparison to decide what needs "
    "re-embedding. Handles four entry points: full workspace index, single file re-index, "
    "file deletion, and file rename."
)

heading3("userId.ts — Stable Identity")
body(
    "Reads the globally configured git email (git config --global user.email) and returns its "
    "MD5 hash. If git email is not configured, the extension shows a clear error and refuses to "
    "index — no silent fallback that would create a different namespace each time."
)

heading2("3.2 — Python Backend Layer")
body(
    "The backend is a plain Python HTTP server using the standard library's http.server.HTTPServer. "
    "No Flask, no FastAPI — only four pip dependencies: openai, pinecone, voyageai, python-dotenv."
)
add_table(
    ["Endpoint", "Method", "Purpose"],
    [
        ["GET /health", "GET", "Liveness check — returns {\"ok\": true}"],
        ["POST /search", "POST", "Normal regex search or AI semantic search"],
        ["POST /index", "POST", "Summarize + embed + upsert new/changed chunks"],
        ["POST /wipe", "POST", "Delete all vectors in a Pinecone namespace"],
    ]
)

heading2("3.3 — External Services")
add_table(
    ["Service", "Model", "Used For", "Cost Model"],
    [
        ["Voyage AI", "voyage-code-2", "Embedding code chunks and queries", "200M tokens/month free"],
        ["OpenAI", "gpt-4o-mini", "Function summarization, line locator", "~$0.00015/1K input tokens"],
        ["Pinecone", "Serverless", "Vector storage and similarity search", "Free tier: 2GB, 100K queries/month"],
    ]
)
page_break()

# ══════════════════════════════════════════════════════ 04 INDEXING ══
heading1("04  Indexing Pipeline — Full Workflow")
body(
    "Indexing converts source code into searchable vectors. It runs on VS Code startup (full "
    "workspace scan), on file save (single file), and via the Re-index Workspace command (full "
    "wipe + rebuild)."
)

diagram(
"""╔═══════════════════════════════════════════════════════════════════╗
║           FULL WORKSPACE INDEXING WORKFLOW                        ║
╚═══════════════════════════════════════════════════════════════════╝

VS Code Opens Workspace
        │
        ▼
Load .smart-search/index.json  (empty {} on first run)
Compute namespace: projectId :: MD5(git email)
        │
        ▼
Walk all files (skip: .git node_modules vendor dist __pycache__ ...)
(keep: .py .ts .tsx .js .jsx .java .go .rs .c .cpp .cs .rb .php .swift .kt)
(skip: files > 500 KB)
        │
        │ for each file:
        ▼
┌──────────────────────────────────────────────────────┐
│ LEVEL 1 — FILE HASH CHECK                            │
│ MD5(entire file content)                             │
│   == stored fileHash?  →  skip file entirely ──────► │ next file
│   ≠  stored fileHash?  →  proceed to Level 2         │
└──────────────────────────────────────────────────────┘
        │
        ▼
Chunk file into functions (TypeScript, local, NO network)
Filter out type="class" (would dominate Pinecone scores)
        │
        │ for each chunk:
        ▼
┌──────────────────────────────────────────────────────┐
│ LEVEL 2 — FUNCTION HASH CHECK                        │
│ MD5(normalizeCode(chunk.content))                    │
│   == stored function hash?  →  keep existing vector  │
│   ≠  stored hash?           →  queue for re-embed    │
│   Not in index?             →  queue as new          │
│   In index but gone?        →  queue for deletion    │
└──────────────────────────────────────────────────────┘
        │
        ▼
POST /index { chunks:[...], delete_ids:[...], namespace }
        │
        ▼  Python backend:
Step 1: delete_chunks(delete_ids, namespace)  [Pinecone]
Step 2: summarize each chunk via GPT          [PARALLEL]
Step 3: build embed text = label+summary+code
Step 4: embed_chunks() — ONE batched API call [Voyage AI]
Step 5: upsert_chunks(vectors, namespace)     [Pinecone]
        │
        ▼
Update index.json with new function hashes
Status bar: "✓ Smart Search: N files updated" """
)

heading2("index.json Structure")
body("The local hash store contains only MD5 hashes — no embeddings, no code content:")
code_block(
"""{
  "src/auth.php": {
    "fileHash": "a3f2c1d4",
    "functions": {
      "verifyToken": {
        "hash":    "d4e2f1b8",
        "chunkId": "src/auth.php::verifyToken"
      },
      "hashPassword": {
        "hash":    "b8c3a2f1",
        "chunkId": "src/auth.php::hashPassword"
      }
    }
  }
}"""
)
body(
    "All paths are relative to the workspace root, making the index portable across machines "
    "and folder renames."
)
page_break()

# ══════════════════════════════════════════════════════ 05 SEARCH ══
heading1("05  Search Pipeline — Normal and AI Modes")

heading2("Normal Search Workflow")
diagram(
"""User types query → vscode.postMessage({ command:"search", searchType:"normal" })
        │
        ▼
POST /search { query, workspacePath, matchCase, matchWord, useRegex,
               filesInclude, filesExclude }
        │
        ▼  Python: search.py
Build regex:
  if useRegex:  pattern = query
  else:         pattern = re.escape(query)
  if matchWord: pattern = \\b{pattern}\\b
  compiled = re.compile(pattern, flags)
        │
        ▼
os.walk(workspacePath):
  skip IGNORE_FOLDERS
  apply filesInclude/filesExclude glob filters
  for each line: compiled.finditer(line)
    → { file, line_num, content, matches:[[start,end],...] }
        │
        ▼
Return { results, total, time_ms }
User clicks result → VS Code opens file, highlights matched chars"""
)

heading2("AI Semantic Search Workflow")
diagram(
"""User types natural language query (≥20 characters)
e.g. "where do we check if subscription has expired"
        │
        ▼
main.js enforces 20-char minimum (shows warning if too short)
        │
        ▼
POST /search { query, searchType:"ai", namespace, threshold:40,
               filesInclude, filesExclude, workspacePath }
        │
        ▼  Python /search (AI branch):
threshold = 40/100 = 0.40

a) embedder.embed_text(query)
   Voyage AI, voyage-code-2, input_type="query"
   → query_vector: [0.21, -0.54, 0.87, ...] (1536 floats)

b) pinecone.query(query_vector, namespace, top_k=10)
   Cosine similarity against all stored function vectors
   → top 10 matches with score + metadata

c) Filter: score >= threshold (0.40)
   Sort: highest score first

d) Apply filesInclude / filesExclude glob filters

e) Line locator — ALL RESULTS IN PARALLEL:
   ThreadPoolExecutor(max_workers=len(results))
   For each result simultaneously:
     1. Read function lines from disk
     2. Build numbered listing: "21: if ($token['exp'] < time()) {"
     3. Ask GPT: which line answers the query?
     4. Parse JSON response, clamp to valid range
     → r["match_line"], r["match_content"]
   All N calls complete in ~one GPT call worth of time (~350ms)

f) Return results with score, summary, match_line, match_content
   User clicks → VS Code opens file at match_line"""
)
page_break()

# ══════════════════════════════════════════════════════ 06 DATA ══
heading1("06  Data Model and Storage")

heading2("Local Storage — .smart-search/")
body("The .smart-search/ folder lives inside the project directory and is auto-added to .gitignore.")
add_table(
    ["File", "Contents", "Size (typical)"],
    [
        ["project-id", "UUID generated on first run, e.g. a3f2c1d4-e5b6-...", "37 bytes"],
        ["index.json", "MD5 hashes of every indexed function (no embeddings)", "~25 KB for 200 functions"],
    ]
)

heading2("Pinecone Record Structure")
add_table(
    ["Field", "Type", "Example"],
    [
        ["id", "string", "src/auth.php::verifyToken"],
        ["values", "float[1536]", "[0.21, -0.54, 0.87, ...]"],
        ["metadata.file", "string", "src/auth.php"],
        ["metadata.name", "string", "verifyToken"],
        ["metadata.type", "string", "function"],
        ["metadata.start_line", "int", "12"],
        ["metadata.end_line", "int", "38"],
        ["metadata.content", "string (≤1000 chars)", "first 1000 chars of function source"],
        ["metadata.summary", "string (≤500 chars)", "GPT-generated plain English description"],
    ]
)
body(
    "The namespace format is {projectId}::{userId}. The double colon (::) is used as separator "
    "because it does not appear in UUIDs (hyphens only) or MD5 hashes (hex only)."
)
page_break()

# ══════════════════════════════════════════════════ 07 CHALLENGES ══
heading1("07  Engineering Challenges and Solutions")
body(
    "This section documents every significant challenge encountered during development and the "
    "solution implemented. These are the real problems that shaped the architecture."
)

# ── C1 ──────────────────────────────────────────────────────────────────────
challenge_header("C-01", "Multi-Language Syntax Parsing")
heading3("The Problem")
body(
    "To index at the function level, the chunker must split any source file into its individual "
    "functions and methods. Every supported language expresses a function differently, and more "
    "critically, every language uses a different mechanism to delimit where a function ends."
)
bullet("Brace languages (PHP, TypeScript, JS, Java, Go, Rust, C#, Swift, Kotlin, C, C++) — blocks start with { and end with the matching }. Counting brace depth is required.")
bullet("Python — no braces. A function body is defined by indentation. A def block ends when a non-blank line at the same or lower indentation appears.")
bullet("Ruby — uses the keyword 'end'. Blocks are opened with def/class/if/do and closed with end. Depth is tracked by counting opener/closer keyword pairs.")
heading3("The Solution: Three Block-End Algorithms")
body("Algorithm 1 — findBraceEnd (all C-style languages):")
code_block(
"""function findBraceEnd(lines, startIdx) {
  let depth = 0, started = false;
  for (let i = startIdx; i < lines.length; i++) {
    for (const ch of lines[i]) {
      if (ch === "{") { depth++; started = true; }
      else if (ch === "}" && started && --depth === 0) return i;
    }
  }
}"""
)
body("Algorithm 2 — findPythonEnd (Python):")
code_block(
"""function findPythonEnd(lines, startIdx) {
  const headerIndent = lines[startIdx].match(/^(\\s*)/)[1].length;
  for (let i = startIdx + 1; i < lines.length; i++) {
    if (lines[i].trim() === "") continue; // blank lines OK inside block
    const indent = lines[i].match(/^(\\s*)/)[1].length;
    if (indent <= headerIndent) return i - 1;
  }
}"""
)
body("Algorithm 3 — findRubyEnd (Ruby):")
code_block(
"""function findRubyEnd(lines, startIdx) {
  let depth = 1;
  const openers = /^\\s*(?:def|class|module|do|if|while|begin|case)\\b/;
  const closer  = /^\\s*end\\b/;
  for (let i = startIdx + 1; i < lines.length; i++) {
    if (openers.test(lines[i])) depth++;
    if (closer.test(lines[i]) && --depth === 0) return i;
  }
}"""
)
callout("NOTE", "Classes are excluded from indexing. A class chunk contains all its methods' code combined, making it always score higher than individual methods in Pinecone — creating noise. Individual methods are already indexed separately.")

# ── C2 ──────────────────────────────────────────────────────────────────────
challenge_header("C-02", "Hash-Based Incremental Indexing (No Re-embed on Every Run)")
heading3("The Problem")
body(
    "The naive approach embeds every function every time VS Code opens. For a 300-function project, "
    "this means 300 GPT calls + 300 Voyage AI calls every startup — ~$0.30 and 15-30 seconds of "
    "waiting. Completely unacceptable for a daily tool."
)
heading3("The Solution: Two-Level Local Hash Store (index.json)")
body("Before any API call, two levels of MD5 comparison happen locally:")
diagram(
"""Level 1 — File hash (one MD5, one JSON lookup per file)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
MD5(entire file content) == stored fileHash?
  YES → skip entire file (no chunking, no API calls)
  NO  → proceed to Level 2

Level 2 — Function hash (one MD5 per function)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
For each function extracted by the chunker:
  MD5(normalize(content)) == stored function hash?
    YES → skip (keep existing Pinecone vector)
    NO  → queue for summarize + embed + upsert
    Not in index → queue as new
    In index but gone from file → queue for deletion

Only the delta (changed + deleted) goes to the backend."""
)
callout("IMPACT", "300-function project, 3 functions changed: 3 GPT calls instead of 300. Saves ~99% of API cost and time on every subsequent run.")

# ── C3 ──────────────────────────────────────────────────────────────────────
challenge_header("C-03", "Formatting-Insensitive Code Normalization")
heading3("The Problem")
body(
    "Auto-formatters (Prettier, Black, gofmt) change whitespace without changing logic. Running "
    "Black on a 100-function Python file would change the MD5 hash of every function — triggering "
    "100 unnecessary re-embeds for changes that mean nothing semantically."
)
heading3("The Solution: normalizeCode()")
code_block(
"""function normalizeCode(content: string): string {
  return content
    .split("\\n")
    .map(line => line.trimEnd())          // remove trailing spaces/tabs
    .filter(line => line.trim() !== "")  // drop blank lines
    .join("\\n")
    .trim();
}"""
)
body("This makes adding/removing blank lines and trailing spaces invisible to the hash comparison. Applied only for comparison — the original code is what gets embedded and displayed.")

# ── C4 ──────────────────────────────────────────────────────────────────────
challenge_header("C-04", "Local Server and the Python Distribution Problem")
heading3("The Problem")
body(
    "The AI features require a Python server. Several scenarios break this: collaborators who "
    "never installed Python, the developer's second laptop, and Windows machines where the "
    "Microsoft Store App Execution Alias intercepts the python command even when Python is installed."
)
heading3("The Solution: Configurable Backend URL")
code_block(
"""// config.ts — ALL backend calls go through this function
export function getBackendUrl(): string {
  const config = vscode.workspace.getConfiguration("smartSearch");
  const url = config.get<string>("backendUrl", "http://localhost:8000");
  return url.replace(/\\/$/, ""); // strip trailing slash
}"""
)
body("Setting smartSearch.backendUrl in VS Code settings allows pointing at a remote server. Collaborators without Python can use the extension by pointing at a shared Python server.")
callout("WINDOWS NOTE", "Disable Microsoft Store App Execution Aliases for python.exe and python3.exe in Settings → Apps → Advanced app settings → App execution aliases.")

# ── C5 ──────────────────────────────────────────────────────────────────────
challenge_header("C-05", "Function Size and the 6 000-Character Cap")
heading3("The Problem")
body(
    "Some functions are very large (400+ lines). Embedding them creates problems: API token limits, "
    "poor embedding quality (a vague average vector), GPT unable to summarize meaningfully, and "
    "wasted API cost."
)
heading3("The Solution: Caps at Each Layer")
add_table(
    ["Layer", "Cap", "Reason"],
    [
        ["chunker.ts", "6 000 chars stored in Chunk", "Limits what is sent to backend"],
        ["summarizer.py", "3 000 chars sent to GPT", "Keeps prompt focused; reduces token cost"],
        ["embedder.py", "8 000 chars sent to Voyage AI", "Near model input limit"],
        ["pinecone_client.py", "1 000 chars for content field", "Pinecone metadata per-field limit"],
        ["pinecone_client.py", "500 chars for summary field", "Keeps result payload compact"],
    ]
)
body(
    "When a function exceeds 6 000 characters, only the first 6 000 are embedded. The function "
    "signature, name, docstring, and opening logic — the most semantically discriminative parts — "
    "are almost always in the first portion."
)

# ── C6 ──────────────────────────────────────────────────────────────────────
challenge_header("C-06", "Relevance Accuracy Threshold")
heading3("The Problem")
body(
    "Pinecone always returns top-K results even if they are poor matches. Showing 10 results "
    "when only 2 are genuinely relevant creates noise and erodes trust in the tool."
)
heading3("The Solution: Configurable Similarity Threshold")
code_block(
"""# server.py — filter and sort by score
results = sorted(
    [r for r in all_results if r["score"] >= threshold],
    key=lambda r: r["score"],
    reverse=True,
)"""
)
add_table(
    ["Threshold", "Behaviour"],
    [
        ["20–30%", "Broad — many results including weak matches. Good for exploration."],
        ["35% (default)", "Balanced — returns genuine matches for clear queries."],
        ["50–60%", "Strict — only very confident matches. May return fewer results."],
        ["70%+", "Very strict — only near-exact semantic matches."],
    ]
)

# ── C7 ──────────────────────────────────────────────────────────────────────
challenge_header("C-07", "Automatic .gitignore Management")
heading3("The Problem")
body(
    "If .smart-search/ is accidentally committed: the project-id gets pushed and collaborators "
    "who clone the repo get the same project-id, potentially writing into the same Pinecone "
    "namespace. The index.json hashes are machine-specific and useless to anyone else."
)
heading3("The Solution: Auto-update .gitignore on Every Run")
code_block(
"""function addToGitignore(workspaceRoot: string): void {
  const entry   = ".smart-search/";
  const comment = "# Smart Search — local function hashes (not committed)";
  const gitignorePath = path.join(workspaceRoot, ".gitignore");

  if (fs.existsSync(gitignorePath)) {
    const content = fs.readFileSync(gitignorePath, "utf8");
    if (!content.includes(entry)) {  // only append if not already there
      fs.appendFileSync(gitignorePath, `\\n${comment}\\n${entry}\\n`);
    }
  } else {
    fs.writeFileSync(gitignorePath, `${comment}\\n${entry}\\n`);
  }
}"""
)
body("Called every time getProjectId() runs (i.e. on every indexing run). Safe to call repeatedly — checks for duplicates before writing.")

# ── C8 ──────────────────────────────────────────────────────────────────────
challenge_header("C-08", "Multi-User Namespace Isolation via Git Email")
heading3("The Problem")
body(
    "If two developers work on the same repo with the same Pinecone account, Developer B's "
    "re-index overwrites Developer A's vectors. Developer A's searches return results shaped "
    "by B's embeddings. Using a random ID would break cross-machine continuity."
)
heading3("The Solution: MD5(git global email) as Stable User ID")
body("Pinecone namespace = {projectId}::{userId} where userId = MD5(git_email.toLowerCase())")
code_block(
"""// userId.ts
email = execSync("git config --global user.email").trim();
// "nawfal@gmail.com"
return crypto.createHash("md5").update(email.toLowerCase()).digest("hex");
// "b7f2a1c5d3e4f6a7..."   """
)
add_table(
    ["Scenario", "git email", "userId", "Outcome"],
    [
        ["Same developer, same machine", "Same", "Same", "Reuses existing vectors"],
        ["Same developer, new machine", "Same (configured by user)", "Same", "Reuses existing vectors"],
        ["VS Code reinstall", "Same (lives in ~/.gitconfig)", "Same", "Nothing changes"],
        ["Friend on same repo", "Different", "Different", "Separate namespace, no conflict"],
        ["Friend tries to edit your vectors", "Different email", "Different namespace", "Cannot touch your vectors"],
    ]
)
body("The email is hashed so the raw address is never sent to Pinecone or any external service. If git email is not configured, the extension shows an error and refuses to index — no silent fallback.")

# ── C9 ──────────────────────────────────────────────────────────────────────
challenge_header("C-09", "Cross-Machine Portability (Same Developer, New Machine)")
heading3("The Problem")
body(
    "When the same developer opens the project on a second machine, .smart-search/ is absent "
    "(gitignored). index.json is empty — every function appears new and would all be re-embedded, "
    "wasting the full initial indexing cost."
)
heading3("How It Works in Practice")
numbered("Developer configures git with the same email on the new machine.")
numbered("Same userId (MD5 of email) → same Pinecone namespace → existing vectors are already there.")
numbered("Empty index.json → all functions queued as new → all re-embedded.")
numbered("Pinecone upsert is idempotent: upserting an existing ID overwrites it, no duplicates.")
numbered("After indexing completes, index.json is populated and future runs are incremental again.")
callout("RESULT", "One-time full re-index cost on a new machine. After that, incremental updates resume. Pinecone vectors are not lost — only the local hash cache needs rebuilding.")

# ── C10 ──────────────────────────────────────────────────────────────────────
challenge_header("C-10", "Why Voyage AI, Not OpenAI Embeddings")
heading3("The Problem")
body(
    "OpenAI's text-embedding-3-small and text-embedding-3-large are trained on general internet "
    "text. They work poorly on code because code uses abbreviations, camelCase names, and "
    "punctuation with semantic meaning ({, :, ->, =>) that general models don't understand well. "
    "Early testing with OpenAI embeddings produced scores of 35–55% for clearly relevant results — "
    "too noisy to set a useful threshold."
)
heading3("The Solution: voyage-code-2")
body("Voyage AI's voyage-code-2 is specifically trained on code + natural language pairs from GitHub, documentation, and technical writing.")
add_table(
    ["Feature", "OpenAI text-embedding-3-small", "Voyage voyage-code-2"],
    [
        ["Training focus", "General internet text", "Code + NL pairs"],
        ["Dimensions", "1536", "1536"],
        ["Asymmetric mode", "No", "Yes (document / query)"],
        ["Typical code score range", "35–55%", "60–85%"],
        ["Free tier", "None", "200M tokens/month"],
    ]
)
heading3("Asymmetric Search — document vs query Modes")
body(
    "Code chunks are embedded with input_type='document' (optimized for retrieval). "
    "Search queries are embedded with input_type='query' (optimized to point toward relevant documents). "
    "This asymmetry means a natural language query vector is positioned in the vector space to be "
    "close to relevant function vectors — even though the query looks nothing like the code syntax."
)

# ── C11 ──────────────────────────────────────────────────────────────────────
challenge_header("C-11", "Minimum Query Length (20 Characters)")
heading3("The Problem")
body(
    "Short queries like 'get user' embed to a vague vector near the center of the semantic space. "
    "Cosine similarity against almost any function is 30–45% — just enough to pass the threshold "
    "but with no genuine signal. The result is a page of results that appear random."
)
heading3("The Solution: 20-Character Minimum in the Frontend")
body("Before sending an AI search request, main.js checks query length. Under 20 characters, a warning is shown and the request is blocked.")
add_table(
    ["Length", "Example", "Status"],
    [
        ["13 chars", "validate token", "Blocked"],
        ["19 chars", "validate auth token", "Blocked"],
        ["26 chars", "validate user auth token", "Allowed"],
        ["45+ chars", "check if the user authentication token has expired", "Ideal"],
    ]
)
body("This constraint only applies to AI search. Normal text/regex search has no minimum length.")

# ── C12 ──────────────────────────────────────────────────────────────────────
challenge_header("C-12", "Line-Level Result Pinpointing")
heading3("The Problem")
body(
    "Pinecone returns results at the function level: 'verifyToken(), lines 5–42'. The developer "
    "must still scan 38 lines to find the specific part answering their query. For a question like "
    "'where do we check token expiry?', the answer is one if statement on line 21."
)
heading3("The Solution: GPT Line Locator")
numbered("Read the function's lines from disk using file path and line numbers from Pinecone metadata.")
numbered("Build a numbered listing: \"21: if ($token['expires_at'] < time()) {\"")
numbered("Send query + numbered listing to gpt-4o-mini, ask for JSON: {\"line\": N, \"content\": \"...\"}")
numbered("Parse response, clamp line number to valid range [start_line, end_line].")
numbered("User clicks result → VS Code opens file at match_line directly.")
code_block(
"""# The prompt sent to GPT:
f'Search query: "{query}"\\n\\n'
f'Code (with line numbers):\\n{numbered[:3000]}\\n\\n'
f'Which single line is most directly relevant? '
f'Reply ONLY with JSON: {{"line": <number>, "content": "<line text>"}}'"""
)
callout("FALLBACK", "If GPT fails or file cannot be read, result falls back to start_line with empty content. Still shown and clickable — just without line-level precision.")

# ── C13 ──────────────────────────────────────────────────────────────────────
challenge_header("C-13", "GPT Summary as Semantic Bridge Between English and Code")
heading3("The Problem")
body(
    "Embedding raw code alone produces a vector that leans toward syntax. A query 'fetch a single "
    "product from the database by its ID' might get 45% similarity against the code that implements "
    "it — below any reasonable threshold — because the code surface is dominated by PHP syntax "
    "($, ->, ?, SELECT) rather than English vocabulary."
)
heading3("The Solution: Pre-Embedding GPT Summarization")
body("GPT generates a plain English summary that is prepended to the code before embedding. The text sent to Voyage AI becomes:")
code_block(
"""Function: getProductById
Fetches a single product's details from the database by its ID using
a SELECT query with a prepared statement. Returns the first row as an
array or false if the product is not found. Includes error logging.

function getProductById($id) {
  $result = $this->db->query('SELECT * FROM products WHERE id = ?', [$id]);
  ...
}"""
)
body("With this bridge, similarity against 'fetch a single product from database by ID' rises from ~45% to ~81% — well above the threshold.")
body("The summary prompt asks GPT to cover: (1) what the function does and why, (2) whether it reads or writes data, (3) whether it contains logging or error handling.")
callout("BONUS", "The summary is stored in Pinecone metadata and displayed in search result cards, giving the developer a plain English description before they click through.")

# ── C14 ──────────────────────────────────────────────────────────────────────
challenge_header("C-14", "Parallel API Calls with ThreadPoolExecutor")
heading3("The Problem")
body(
    "Two pipeline stages require multiple LLM calls: summarizing N new functions (one GPT call each) "
    "and running the line locator on up to 10 results (one GPT call each). Sequential execution: "
    "10 summaries = ~3 500ms. 8 line-locates = ~2 800ms. Both unacceptable."
)
heading3("The Solution: ThreadPoolExecutor")
body("All LLM calls within a batch are submitted simultaneously. Since GPT calls are I/O-bound (waiting for HTTP responses), Python's GIL does not limit parallelism — network I/O releases the GIL.")
code_block(
"""# Parallel summarization of N chunks
with ThreadPoolExecutor(max_workers=max(1, len(chunks_to_embed))) as executor:
    futures = {
        executor.submit(summarize_chunk, chunk): chunk
        for chunk in chunks_to_embed  # all submitted simultaneously
    }
    for future in as_completed(futures):
        chunk = futures[future]
        try:
            chunk["summary"] = future.result()
        except Exception:
            chunk["summary"] = ""  # non-fatal: embed without summary"""
)
add_table(
    ["Operation", "N items", "Sequential", "Parallel", "Speedup"],
    [
        ["Summarize functions", "10", "~3 500ms", "~400ms", "8.75×"],
        ["Line-locate results", "8", "~2 800ms", "~350ms", "8×"],
        ["Summarize functions", "3", "~1 050ms", "~400ms", "2.6×"],
    ]
)
callout("RESULT", "Total AI search time (embed query + Pinecone + line-locate 8 results) consistently under 1 second. Indexing 10 new functions after a file save completes in under 500ms.")
page_break()

# ══════════════════════════════════════════════════ 08 WORKED EXAMPLE ══
heading1("08  End-to-End Worked Example")

heading2("Scenario A: First-Time Index on a New Project")
diagram(
"""1. VS Code opens /projects/ecommerce-api

2. extension.ts activate():
   getProjectId(): creates .smart-search/, generates UUID
   getUserId():    git config --global user.email → MD5
   namespace = "a3f2c1d4-...::b7f2a1c5..."
   loadIndex() → {} (empty)
   walkFiles() → 47 files found

3. All 47 files: no stored hashes → chunk all → queue all 183 functions

4. POST /index { chunks:[183], delete_ids:[], namespace }

5. Python:
   ThreadPoolExecutor: 183 summarize_chunk() calls in parallel
   Time: ~500ms (vs ~64 seconds sequential)
   embed_chunks(): ONE batched Voyage AI call → 183 vectors
   upsert_chunks(): 183 records written to Pinecone

6. Update index.json with 183 function hashes
   Status bar: "✓ Smart Search: 47 files updated"
   Total time: ~3 seconds"""
)

heading2("Scenario B: AI Search — \"where do we check subscription expiry\"")
diagram(
"""1. User types: "where do we check if subscription has expired" (45 chars)
2. Clicks AI search, threshold=40

3. POST /search { searchType:"ai", namespace, threshold:40 }

4. Python:
   embed_text(query) → query_vector  [Voyage AI, ~100ms]
   pinecone.query(query_vector, namespace, top_k=10)  [~50ms]

   Results:
   #1 score=0.84 → SubscriptionService.php::checkActive
   #2 score=0.79 → SubscriptionService.php::isExpired
   #3 score=0.61 → AuthMiddleware.php::validateToken
   #4 score=0.42 → Invoice.php::getDueDate
   #5 score=0.38 → (filtered — below 0.40 threshold)

   Line locator — 4 results in PARALLEL:
   Thread 1 → line 28: "if ($this->expires_at < now()) {"
   Thread 2 → line 67: ...
   Thread 3, 4: ...
   All complete in ~380ms

5. User sees 4 result cards with score badges and summaries
6. Clicks result #1 → VS Code opens SubscriptionService.php line 28"""
)
page_break()

# ══════════════════════════════════════════════════ 09 PERFORMANCE ══
heading1("09  Performance Analysis")

heading2("Indexing Performance")
add_table(
    ["Scenario", "Files", "Functions", "Time", "API Cost (est.)"],
    [
        ["First run (cold index)", "47", "183", "~3s", "~$0.005"],
        ["Subsequent open (0 changes)", "47", "0 changed", "<100ms", "$0.00"],
        ["File save (3 funcs changed)", "1", "3", "~500ms", "<$0.001"],
        ["Re-index workspace command", "47", "183", "~3s", "~$0.005"],
    ]
)

heading2("Search Performance")
add_table(
    ["Operation", "Time", "Notes"],
    [
        ["Normal search (text/regex)", "50–300ms", "Depends on project size and disk speed"],
        ["AI: query embedding (Voyage AI)", "~100ms", "Single API call"],
        ["AI: Pinecone query (top_k=10)", "~50ms", "Serverless tier"],
        ["AI: line locator (8 results, parallel)", "~350ms", "8 simultaneous gpt-4o-mini calls"],
        ["AI: total end-to-end", "~600–900ms", "Under 1 second consistently"],
    ]
)
page_break()

# ══════════════════════════════════════════════════ 10 SECURITY ══
heading1("10  Security and Privacy Design")

heading2("Source Code Exposure")
add_table(
    ["Service", "What is sent", "When"],
    [
        ["OpenAI (summarizer)", "First 3 000 chars of each new/changed function", "At index time, once per function change"],
        ["Voyage AI (embedder)", "First 8 000 chars of summary+code", "At index time, once per function change"],
        ["Pinecone", "Vectors (floats) + max 1 000 chars content metadata", "At index time; full source never stored"],
    ]
)

heading2("Key Security Decisions")
bullet("All API keys live in backend/.env which is gitignored. Never committed or shared.")
bullet("Git email is hashed with MD5 before use. Raw email is never sent anywhere external.")
bullet("index.json contains only MD5 hashes and relative file paths — no code, no keys, no embeddings.")
bullet("Backend binds to localhost by default. Not exposed to the internet unless explicitly configured.")
page_break()

# ══════════════════════════════════════════════════ 11 FUTURE WORK ══
heading1("11  Future Work")

heading2("1. Sliding-Window Chunking for Very Large Functions")
body(
    "Functions longer than 6 000 characters are currently truncated. A better approach is to split "
    "them into overlapping windows (e.g. 3 000-char windows with 500-char overlap) and index each "
    "window as a separate chunk. Queries would then match any part of a large function."
)

heading2("2. Multi-Root Workspace Support")
body(
    "VS Code supports opening multiple folders simultaneously. The extension currently processes only "
    "workspaceFolders[0]. Supporting all roots would require per-root namespacing and status tracking."
)

heading2("3. Cross-Project Search")
body(
    "A 'search all my projects' mode could query multiple Pinecone namespaces and merge results. "
    "Useful for developers maintaining multiple microservices."
)

heading2("4. Query Result Caching")
body(
    "Repeated identical queries within a session currently hit the full pipeline. A short TTL "
    "in-memory cache would make repeated searches instant."
)

heading2("5. Streaming Search Results")
body(
    "Currently the UI shows a spinner until all line-locator calls complete. Streaming partial "
    "results as each locator call finishes would make the UI feel more responsive."
)

heading2("6. Automatic Model Version Upgrades")
body(
    "Voyage AI periodically releases improved models. Vectors stored with an older model cannot "
    "be queried by a newer model. An upgrade path would detect which model was used and trigger "
    "a re-index when the configured model changes."
)

heading2("7. Support for More File Types")
body(
    "Adding SQL stored procedures, Solidity smart contracts, Dart (Flutter), Scala, Haskell, etc. "
    "The chunker architecture is designed for easy extension — adding a new language is a "
    "self-contained function plus a case in the switch statement."
)
page_break()

# ══════════════════════════════════════════════════ 12 CONCLUSION ══
heading1("12  Conclusion")
body(
    "Smart Search demonstrates that production-quality semantic code search can be built as a solo "
    "personal tool, using cloud AI services composed around a VS Code extension. The key insight "
    "driving the whole system is that finding code by meaning rather than by keyword requires three "
    "technical bridges: a code-specialized embedding model (Voyage AI voyage-code-2), a plain English "
    "semantic bridge (GPT-generated summaries), and line-level precision (a second GPT call that "
    "pinpoints the exact line within a matched function)."
)
body(
    "The engineering challenges were not in the AI itself — the models work well once correctly "
    "integrated — but in the surrounding systems: keeping re-indexing cheap (hash-based change "
    "detection), handling multi-user isolation correctly (git email namespacing), parsing 11 "
    "different language syntaxes reliably (three block-end algorithms), making parallel LLM calls "
    "fast enough for a developer tool (ThreadPoolExecutor), and making the system portable across "
    "machines and collaborators (relative paths, configurable backend URL)."
)
body(
    "Each challenge was solved with a targeted, minimal solution. The overall architecture is simple "
    "enough to understand in an afternoon but robust enough to handle the daily workflow of a real "
    "development project. The result is a tool that meaningfully changes how code navigation works — "
    "where the question 'where do we validate payment subscription status?' returns a specific file, "
    "function, and line number in under one second, without requiring any knowledge of what that "
    "function is called."
)

divider()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Smart Search  ·  Technical Research Report  ·  Nawfal Jalloul  ·  May 2026\n"
    "VS Code Extension API  ·  Voyage AI voyage-code-2  ·  OpenAI gpt-4o-mini  ·  Pinecone Serverless"
)
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.save("report.docx")
print("report.docx created successfully.")
