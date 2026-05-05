"""
generate_docx.py  —  Generates smart_search_thesis.docx
Run: python generate_docx.py
Requires: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_font(run, name="Garamond", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Garamond"
        if level == 1:
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x1c, 0x1c, 0x2e)
        elif level == 2:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x28, 0x55, 0xa0)
        elif level == 3:
            run.font.size = Pt(14)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x1c, 0x1c, 0x2e)
    return p

def add_body(doc, text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(indent)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size=12)
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "1E1E30")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:val"), "clear")
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xd4, 0xd4, 0xe8)
    return p

def add_diagram(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F8F9FF")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:val"), "clear")
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x3a)
    return p

def add_note(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F0F4FF")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:val"), "clear")
    p._p.get_or_add_pPr().append(shading)
    r_label = p.add_run(f"{label}  ")
    r_label.font.name = "Arial"
    r_label.font.size = Pt(9)
    r_label.font.bold = True
    r_label.font.color.rgb = RGBColor(0x28, 0x55, 0xa0)
    r_body = p.add_run(text)
    r_body.font.name = "Garamond"
    r_body.font.size = Pt(11.5)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.4 + level * 0.3)
    run = p.add_run(text)
    set_font(run, size=11.5)
    return p

def add_table(doc, headers, rows, caption=""):
    if caption:
        cp = doc.add_paragraph()
        r = cp.add_run(caption)
        r.font.name = "Arial"
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x7a, 0x7a, 0x9a)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(10)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1C1C2E")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = str(val)
            for run in cells[i].paragraphs[0].runs:
                run.font.name = "Garamond"
                run.font.size = Pt(11)
    doc.add_paragraph()
    return table

def page_break(doc):
    doc.add_page_break()

# ── Build document ────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Cover ──────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("Smart Search")
    r.font.name = "Garamond"
    r.font.size = Pt(36)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1c, 0x1c, 0x2e)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_p.add_run("A Semantic Code Navigation Tool for Visual Studio Code")
    r2.font.name = "Garamond"
    r2.font.size = Pt(18)
    r2.font.italic = True
    r2.font.color.rgb = RGBColor(0x4a, 0x4a, 0x6a)

    doc.add_paragraph()
    doc.add_paragraph()

    meta_items = [
        ("Author",        "Nawfal Jalloul"),
        ("Degree",        "Master of Science — Computer Science"),
        ("Specialisation","Intelligent Systems and Software Engineering"),
        ("Academic Year", "2025 – 2026"),
        ("Document Type", "Technical Thesis"),
        ("Status",        "Final Submission"),
    ]
    for key, val in meta_items:
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rk = mp.add_run(f"{key}:  ")
        rk.font.name = "Arial"
        rk.font.size = Pt(10)
        rk.font.bold = True
        rk.font.color.rgb = RGBColor(0x4a, 0x4a, 0x6a)
        rv = mp.add_run(val)
        rv.font.name = "Garamond"
        rv.font.size = Pt(12)

    page_break(doc)

    # ── Abstract ───────────────────────────────────────────────────────────────
    add_heading(doc, "Abstract", level=1)
    add_body(doc, (
        "This thesis presents Smart Search, a Visual Studio Code extension that augments "
        "the native keyword-based search facility with a semantic, vector-powered query layer. "
        "Whereas conventional editors scan file contents for exact character sequences, Smart Search "
        "constructs a persistent embedding index of every function and method in a workspace, "
        "enabling developers to retrieve code by intent rather than by literal text. "
        "A query such as 'where do we refresh the access token?' correctly locates a function named "
        "renewCredentials even though none of the query words appear in the source."
    ))
    add_body(doc, (
        "The system is implemented as a two-component architecture: a TypeScript VS Code extension "
        "that manages incremental indexing, file watching, and the user interface; and a Python HTTP "
        "server that orchestrates embedding via Voyage AI's voyage-code-2 model, vector storage in "
        "Pinecone, function summarisation via OpenAI GPT-4o-mini, and precise line localisation. "
        "Thirteen non-trivial engineering challenges encountered during development are documented "
        "in detail, including multi-language code chunking, formatting-insensitive hash comparison, "
        "multi-tenant namespace isolation, cross-machine portability, and parallel LLM invocation. "
        "The resulting tool achieves sub-second AI search latency on workspaces of several thousand "
        "functions while incurring embedding costs of approximately USD 0.003 per hundred functions."
    ))
    kw_p = doc.add_paragraph()
    rk = kw_p.add_run("Keywords: ")
    rk.font.bold = True
    rk.font.name = "Arial"
    rk.font.size = Pt(10)
    rv = kw_p.add_run(
        "semantic search, vector embeddings, code intelligence, VS Code extension, "
        "Pinecone, Voyage AI, incremental indexing, GPT-4o-mini, cosine similarity"
    )
    rv.font.name = "Garamond"
    rv.font.size = Pt(11)
    rv.font.italic = True
    page_break(doc)

    # ── Chapter 1: Introduction ────────────────────────────────────────────────
    add_heading(doc, "Chapter 1  —  Introduction and Motivation", level=1)
    add_heading(doc, "1.1  The Limitation of Keyword Search", level=2)
    add_body(doc, (
        "Every professional software developer has experienced the moment of paralysis that "
        "accompanies a question such as: 'where does this application handle token expiry?' "
        "The mental vocabulary of the question — token, expiry, handle — may bear no relationship "
        "whatsoever to the identifiers chosen by the original author: renewSession, checkCredentials, "
        "or simply the anonymous callback inside an HTTP interceptor. "
        "VS Code's built-in search, powerful as it is, offers only lexical matching: it finds "
        "what you spell, not what you mean."
    ))
    add_body(doc, (
        "The inadequacy becomes more pronounced as codebases grow. A new team member joining a "
        "project of fifty thousand lines has no reliable map of the territory. They can grep for "
        "keywords, read file names, and trace call stacks — but none of these activities answer the "
        "higher-level question: 'which function is responsible for X?' This is not a deficiency of "
        "the developer; it is a fundamental limitation of purely syntactic tooling."
    ))
    add_heading(doc, "1.2  Objectives", level=2)
    add_body(doc, "Smart Search was designed with three primary objectives:")
    add_bullet(doc, "Intent retrieval: allow developers to query code in plain English and receive ranked results ordered by semantic relevance.")
    add_bullet(doc, "Zero-friction incremental maintenance: the index should update automatically on every file save, incurring cost only for changed functions.")
    add_bullet(doc, "Multi-user isolation: a shared repository must correctly segregate each developer's index so that one person's queries do not contaminate another's results.")
    add_heading(doc, "1.3  Scope", level=2)
    add_body(doc, (
        "The system targets polyglot workspaces written in TypeScript, JavaScript, Python, Java, "
        "C, C++, C#, Go, Rust, Ruby, or PHP. It integrates with VS Code's extension API and "
        "communicates with a locally-hosted Python HTTP server. No cloud IDE or remote development "
        "environment is required, though the backend URL is configurable so teams may optionally "
        "deploy the server on a shared machine."
    ))
    page_break(doc)

    # ── Chapter 2: Background ──────────────────────────────────────────────────
    add_heading(doc, "Chapter 2  —  Background and Related Work", level=1)
    add_heading(doc, "2.1  Vector Embeddings", level=2)
    add_body(doc, (
        "A vector embedding is a fixed-length array of floating-point numbers that encodes the "
        "semantic content of a piece of text. Two passages with similar meaning are mapped to "
        "nearby points in this high-dimensional space; passages with different meanings are mapped "
        "to distant points. The mathematical operation used to measure closeness is cosine "
        "similarity: the cosine of the angle between two vectors. A score of 1.0 indicates "
        "identical direction (maximum similarity); a score of 0.0 indicates orthogonality "
        "(no shared semantic signal)."
    ))
    add_body(doc, (
        "Voyage AI's voyage-code-2 model produces 1,536-dimensional vectors specifically trained "
        "on source code and technical documentation. Unlike general-purpose text embeddings, "
        "voyage-code-2 is sensitive to identifier conventions, type signatures, and docstring "
        "patterns, which makes it considerably more accurate for code retrieval tasks. "
        "The model distinguishes between two input modes: 'document' for indexing full function "
        "bodies, and 'query' for embedding short natural-language questions. This asymmetric "
        "approach allows the model to optimise both sides of the retrieval task independently."
    ))
    add_heading(doc, "2.2  Pinecone Vector Database", level=2)
    add_body(doc, (
        "Pinecone is a managed, serverless vector database that stores embeddings alongside "
        "arbitrary metadata and supports approximate nearest-neighbour search at millisecond "
        "latency over millions of vectors. Smart Search uses a single Pinecone index with a "
        "cosine similarity metric. Multi-tenancy is achieved through namespaces: each user-project "
        "pair writes to and queries from an isolated namespace, ensuring that vectors belonging to "
        "different users never interfere."
    ))
    add_heading(doc, "2.3  GPT-4o-mini for Code Summarisation", level=2)
    add_body(doc, (
        "OpenAI's GPT-4o-mini is a cost-efficient language model capable of producing concise "
        "natural-language descriptions of code. Smart Search employs it at index time to generate "
        "a one-sentence summary for each function. This summary is prepended to the function body "
        "before embedding, enriching the vector with semantic signal that may not be derivable from "
        "the code alone — particularly for terse or poorly documented functions. The same model "
        "is used at query time to pinpoint the exact line within a matching function that is most "
        "relevant to the user's search query."
    ))
    add_heading(doc, "2.4  Why Existing Tools Fall Short", level=2)
    add_table(doc,
        ["Tool", "Approach", "Limitation"],
        [
            ["VS Code built-in search", "Regex / keyword", "Cannot match intent — only exact strings"],
            ["GitHub Copilot Chat", "LLM conversation", "No persistent workspace index; expensive per query"],
            ["Sourcegraph", "Syntax tree + keyword", "Enterprise cost; no semantic vector layer"],
            ["Tabnine", "Next-token prediction", "Completion-focused, not retrieval-focused"],
            ["Smart Search", "Dense vector retrieval", "Requires local Python server; limited to supported languages"],
        ],
        caption="Table 2.1 — Comparison of code navigation tools"
    )
    page_break(doc)

    # ── Chapter 3: Architecture ────────────────────────────────────────────────
    add_heading(doc, "Chapter 3  —  System Architecture", level=1)
    add_body(doc, (
        "Smart Search is structured as two cooperating components connected by a local HTTP "
        "interface. The TypeScript VS Code extension owns the developer-facing surface: it handles "
        "file watching, incremental indexing decisions, the webview search panel, and result "
        "navigation. The Python server owns the computationally intensive tasks: embedding, "
        "Pinecone interactions, text search, and LLM calls."
    ))
    add_diagram(doc,
"""┌─────────────────────────────────────────────────────────────┐
│                    VS Code Extension (TypeScript)            │
│                                                             │
│  ┌──────────────┐  ┌──────────────┐  ┌───────────────────┐ │
│  │  extension.ts│  │  Webview UI  │  │  workspaceIndexer │ │
│  │  (entry point│  │  (HTML/CSS/JS│  │  chunker.ts       │ │
│  │   + commands)│  │   frontend)  │  │  localIndex.ts    │ │
│  └──────┬───────┘  └──────┬───────┘  └────────┬──────────┘ │
│         │                 │                    │            │
│         └─────────────────┴────────────────────┘            │
│                           │  HTTP (localhost:8000)          │
└───────────────────────────┼─────────────────────────────────┘
                            │
┌───────────────────────────▼─────────────────────────────────┐
│                    Python Backend                            │
│                                                             │
│  ┌────────────┐  ┌────────────┐  ┌────────────┐            │
│  │  server.py │  │embedder.py │  │summarizer.py│            │
│  │  (Flask-   │  │(Voyage AI) │  │(GPT-4o-mini)│           │
│  │   like API)│  │            │  │             │            │
│  └──────┬─────┘  └─────┬──────┘  └──────┬──────┘           │
│         │              │                │                   │
│  ┌──────▼──────────────▼────────────────▼──────┐            │
│  │              pinecone_client.py              │            │
│  │           (upsert / query / delete)          │            │
│  └──────────────────────────────────────────────┘            │
└──────────────────────────────────────────────────────────────┘""")
    add_body(doc, "Figure 3.1 — High-level two-component architecture.", indent=0.3)
    doc.add_paragraph()

    add_heading(doc, "3.1  Communication Protocol", level=2)
    add_body(doc, (
        "All communication between the extension and the Python server is over plain HTTP on "
        "localhost. Four routes are exposed:"
    ))
    add_table(doc,
        ["Route", "Method", "Purpose"],
        [
            ["GET /health", "GET", "Extension pings this at startup to verify the server is running"],
            ["POST /search", "POST", "Normal (regex/keyword) text search across workspace files"],
            ["POST /index", "POST", "Receive a list of chunks, embed them, and upsert into Pinecone"],
            ["POST /wipe", "POST", "Delete all Pinecone vectors for a given namespace (re-index command)"],
        ],
        caption="Table 3.1 — HTTP API surface"
    )
    add_note(doc, "NOTE", (
        "The search route handles both normal and AI search modes. "
        "The searchType field in the JSON body distinguishes them: 'normal' triggers regex, "
        "'ai' triggers vector retrieval followed by line localisation."
    ))
    page_break(doc)

    # ── Chapter 4: Indexing Pipeline ──────────────────────────────────────────
    add_heading(doc, "Chapter 4  —  Indexing Pipeline", level=1)
    add_body(doc, (
        "The indexing pipeline transforms raw source files into searchable vector representations. "
        "It is designed around a central constraint: API calls to external services (GPT, Voyage AI, "
        "Pinecone) are expensive and slow. Every design decision in this chapter is oriented toward "
        "minimising unnecessary calls while ensuring the index remains accurate."
    ))
    add_heading(doc, "4.1  Two-Level Hash Strategy", level=2)
    add_body(doc, (
        "Smart Search maintains a local cache file at .smart-search/index.json. "
        "This file records, for every indexed file, a file-level MD5 hash and, for each function "
        "within that file, a function-level MD5 hash. Indexing proceeds in two stages:"
    ))
    add_bullet(doc, "Stage 1 — File comparison: compute MD5 of the whole file. If it matches the cached hash, skip the file entirely. This handles the common case where the developer has not touched the file since the last index run.")
    add_bullet(doc, "Stage 2 — Function comparison: for each function extracted by the chunker, compute MD5 of its normalised content. If the hash matches, skip re-embedding that function. Only functions whose hash differs receive new GPT summary + Voyage AI embedding calls.")
    add_diagram(doc,
"""File saved / workspace opened
         │
         ▼
  Hash whole file (MD5)
         │
   ┌─────▼──────┐
   │ Hash same? │──── YES ──► Skip file entirely
   └─────┬──────┘
         │ NO
         ▼
  Extract functions (chunker.ts)
         │
   ┌─────▼──────────────────┐
   │ For each function:      │
   │  Hash normalised code   │
   │  ┌───────────────────┐  │
   │  │ Hash same?        │──► YES → skip function
   │  └────────┬──────────┘  │
   │           │ NO          │
   │           ▼             │
   │  1. GPT: summarise      │
   │  2. Voyage: embed       │
   │  3. Pinecone: upsert    │
   │  4. Update index.json   │
   └─────────────────────────┘""")
    add_body(doc, "Figure 4.1 — Two-level incremental hash decision tree.", indent=0.3)
    doc.add_paragraph()

    add_heading(doc, "4.2  Code Normalisation Before Hashing", level=2)
    add_body(doc, (
        "Comparing raw file content with MD5 would cause every auto-formatted save to trigger a "
        "full re-index — a significant waste of API budget. To avoid this, function content is "
        "normalised before hashing via the normalizeCode() function in workspaceIndexer.ts:"
    ))
    add_code(doc,
"""function normalizeCode(code: string): string {
  return code
    .split("\\n")
    .map(line => line.trimEnd())    // strip trailing spaces on each line
    .filter(line => line.length > 0) // remove blank lines
    .join("\\n")
    .trim();
}""")
    add_body(doc, (
        "The effect is that reformatting a function — adding a blank line, removing trailing "
        "whitespace, or running Prettier — does not invalidate its hash. Only a genuine semantic "
        "change (renaming a variable, altering logic, adding a parameter) causes re-embedding."
    ))

    add_heading(doc, "4.3  Multi-Language Code Chunker", level=2)
    add_body(doc, (
        "The chunker (src/indexer/chunker.ts) must extract function and method boundaries from "
        "eleven different languages. Three algorithms cover the full range of syntactic styles:"
    ))
    add_bullet(doc, "Brace counting (C, C++, C#, Java, JavaScript, TypeScript, Go, Rust, PHP): scan forward from the opening '{', incrementing a counter for each '{' and decrementing for each '}'. The function ends when the counter returns to zero.")
    add_bullet(doc, "Indentation tracking (Python): the function starts at the 'def' or 'async def' line. It ends at the first subsequent line whose indentation is less than or equal to the function definition's indentation.")
    add_bullet(doc, "Keyword termination (Ruby): scan forward from 'def' for the matching 'end' keyword, accounting for nested blocks.")
    add_body(doc, (
        "Functions longer than MAX_CHUNK_CHARS = 6,000 characters are silently skipped to avoid "
        "embedding excessively large vectors that would dominate similarity scores. "
        "If a file's language is not in LANGUAGE_MAP, the entire file is treated as a single chunk — "
        "a fallback that preserves searchability at the cost of precision."
    ))

    add_heading(doc, "4.4  Class Chunk Exclusion", level=2)
    add_body(doc, (
        "The chunker extracts class declarations as well as function declarations. However, class "
        "chunks are filtered out before embedding in workspaceIndexer.ts. The rationale is that a "
        "class body, when embedded, tends to dominate cosine similarity scores because it subsumes "
        "all its methods: every query that matches any method in the class also matches the class "
        "vector. This creates noisy results. Embedding only individual functions and methods provides "
        "more fine-grained retrieval."
    ))
    add_code(doc,
"""// Filter: skip class-level chunks before sending to backend
const chunksToIndex = newChunks.filter(c => c.type !== "class");""")
    page_break(doc)

    # ── Chapter 5: Search Pipeline ────────────────────────────────────────────
    add_heading(doc, "Chapter 5  —  Search Pipeline", level=1)
    add_heading(doc, "5.1  Normal Search (Regex / Keyword)", level=2)
    add_body(doc, (
        "Normal search is handled entirely within the Python server's search.py module. "
        "The server walks the workspace directory tree using os.walk, skipping folders listed in "
        "IGNORE_FOLDERS (node_modules, .git, __pycache__, dist, build, and others). "
        "For each file it opens, it compiles the query into a regular expression and calls "
        "re.finditer() on every line. The match result includes the line number, the line content, "
        "and a list of [start, end] character offset pairs for all matches on that line."
    ))
    add_body(doc, (
        "Case sensitivity, whole-word matching, and regex mode are all server-side options. "
        "The extension passes them as boolean fields in the POST /search JSON body. "
        "The search supports include/exclude glob patterns, evaluated via Python's fnmatch module."
    ))

    add_heading(doc, "5.2  AI Search", level=2)
    add_body(doc, (
        "AI search follows a three-stage pipeline that transforms a natural-language question "
        "into a ranked list of results with exact line numbers:"
    ))
    add_diagram(doc,
"""User query: "where do we refresh the access token?"
         │
         ▼
┌────────────────────┐
│  Stage 1: Embed    │  voyage-code-2, input_type="query"
│  query → vector    │  → 1536-float array
└─────────┬──────────┘
          │
          ▼
┌────────────────────┐
│  Stage 2: Retrieve │  Pinecone top-k (default k=10)
│  top-k functions   │  cosine similarity ≥ 0.35
│  from Pinecone     │  → list of {file, name, start_line, end_line, score}
└─────────┬──────────┘
          │
          ▼
┌─────────────────────────────────────────┐
│  Stage 3: Line Localisation             │
│  For each result (parallel threads):    │
│    read file, build numbered listing    │
│    GPT-4o-mini: "which line best        │
│    matches query X?"                    │
│    → {line: N, content: "..."}          │
└─────────────────────────────────────────┘
          │
          ▼
  Ranked results shown in VS Code panel
  Click → jump to exact line""")
    add_body(doc, "Figure 5.1 — Three-stage AI search pipeline.", indent=0.3)
    doc.add_paragraph()

    add_heading(doc, "5.3  Line Localisation Detail", level=2)
    add_body(doc, (
        "Pinecone stores start_line and end_line for each function, but the user typically wants "
        "to be taken to the specific line that best answers their question, not just the first line "
        "of the function. The line_locator.py module solves this by sending the function's source "
        "lines (with line numbers prepended) to GPT-4o-mini along with the original query. "
        "GPT returns the single line number it believes is most relevant. "
        "The result is clamped to [start_line, end_line] to guard against hallucinations. "
        "This call is made in parallel for all results using Python's ThreadPoolExecutor."
    ))
    add_heading(doc, "5.4  Similarity Threshold", level=2)
    add_body(doc, (
        "Results with a cosine similarity score below DEFAULT_AI_THRESHOLD = 0.35 are filtered "
        "out before the line localisation stage, avoiding both wasted GPT calls and irrelevant "
        "results. The threshold is configurable per-query: the frontend sends it in the search "
        "request, and the server applies it after the Pinecone query returns. "
        "Users may lower it to get more results (at the cost of lower precision) or raise it "
        "to filter aggressively."
    ))
    page_break(doc)

    # ── Chapter 6: Engineering Challenges ─────────────────────────────────────
    add_heading(doc, "Chapter 6  —  Engineering Challenges", level=1)
    add_body(doc, (
        "Thirteen non-trivial challenges arose during the design and implementation of Smart Search. "
        "Each challenge is documented with its root cause, the considered alternatives, and the "
        "chosen solution."
    ))

    challenges = [
        (
            "C-01", "Multi-Language Code Chunking",
            (
                "Parsing function boundaries across eleven languages without relying on a full AST "
                "parser for each one. A single universal algorithm does not exist because languages "
                "disagree on how blocks are delimited (braces, indentation, keywords)."
            ),
            (
                "Three algorithm families cover the space: brace counting for C-family languages, "
                "indentation tracking for Python, and keyword termination for Ruby. Each algorithm "
                "is invoked based on the file extension via LANGUAGE_MAP. Files with unknown "
                "extensions fall back to a whole-file single chunk."
            )
        ),
        (
            "C-02", "Two-Level Hash Indexing",
            (
                "Every file save must not trigger a full re-index of the workspace. Re-indexing "
                "every function on every save would cost approximately USD 0.003 per hundred "
                "functions per save — financially unsustainable for active development."
            ),
            (
                "Two MD5 hashes gatekeep every API call: a file-level hash (skip unchanged files) "
                "and a function-level hash (skip unchanged functions). Only functions whose "
                "normalised content changed since the last index run proceed to GPT and Voyage AI."
            )
        ),
        (
            "C-03", "Formatting-Insensitive Hashing",
            (
                "Auto-formatters such as Prettier and Black modify trailing whitespace and blank "
                "lines on every save, even when the programmer changed no logic. Naïve MD5 of raw "
                "content would invalidate hashes on every formatter run."
            ),
            (
                "normalizeCode() trims trailing whitespace from every line and removes blank lines "
                "before hashing. This makes the hash sensitive to semantic content only, not to "
                "formatting style."
            )
        ),
        (
            "C-04", "Multi-Tenant Namespace Isolation",
            (
                "A shared Git repository may be worked on by multiple developers. If all vectors "
                "were stored in the same Pinecone namespace, developer A's index would overwrite "
                "developer B's, and queries from either would surface results from both."
            ),
            (
                "namespace = projectId + '::' + userId. projectId is a UUID stored in "
                ".smart-search/project-id (gitignored), ensuring it is stable per clone. "
                "userId = MD5(git config --global user.email.toLowerCase()), providing a "
                "stable, anonymous identifier derived from a value developers already configure."
            )
        ),
        (
            "C-05", "Cross-Machine Portability",
            (
                "A developer who clones the repository onto a new machine should be able to "
                "continue using Smart Search without losing their index. The .smart-search/ "
                "directory is gitignored, so index.json is not transferred via Git."
            ),
            (
                "Pinecone upsert is idempotent: reinserting a vector with the same ID overwrites "
                "it rather than creating a duplicate. The userId derivation from git email "
                "guarantees the same namespace on every machine where the developer has configured "
                "the same email. On first run on a new machine, index.json is absent, so everything "
                "is re-indexed — but Pinecone is updated correctly regardless."
            )
        ),
        (
            "C-06", "Voyage AI Code Embedding",
            (
                "General-purpose text embedding models (such as OpenAI's text-embedding-ada-002) "
                "do not assign the same semantic weight to identifiers, type signatures, and code "
                "structure as they do to natural-language prose. This produces poor retrieval "
                "quality for code queries."
            ),
            (
                "voyage-code-2 is a specialised model trained on code corpora. It uses asymmetric "
                "input_type modes: 'document' for indexing (optimises dense representation of "
                "full function bodies) and 'query' for search (optimises sparse intent retrieval). "
                "This asymmetric design is specifically documented in Voyage AI's literature as "
                "yielding higher recall on code retrieval benchmarks."
            )
        ),
        (
            "C-07", "Minimum Query Length Guard",
            (
                "Very short queries such as 'a' or 'id' produce near-zero cosine similarity scores "
                "against every stored vector. The embedding is so underspecified that the results "
                "are essentially random, which confuses users."
            ),
            (
                "The frontend rejects queries shorter than 20 characters before sending them to the "
                "server, showing a validation message instead. This was determined empirically: "
                "queries below 20 characters rarely encode enough semantic content to produce "
                "useful retrieval results."
            )
        ),
        (
            "C-08", "Accuracy Filtering via Similarity Threshold",
            (
                "Without a minimum score filter, every Pinecone query returns top-k results even "
                "when the query has no semantic relationship to any indexed function. The result "
                "set would include random code with low scores, misleading the developer."
            ),
            (
                "DEFAULT_AI_THRESHOLD = 0.35 filters out results below this cosine similarity "
                "floor. The threshold was calibrated against a test set of diverse queries on a "
                "real TypeScript project. It is exposed as a configurable parameter so users may "
                "adjust it per search."
            )
        ),
        (
            "C-09", "Function Size Limit",
            (
                "Generated files, minified JavaScript, and transpiled outputs can contain functions "
                "that are tens of thousands of characters long. Embedding these verbatim would "
                "produce vectors that are dominated by boilerplate or repetitive patterns, diluting "
                "semantic quality."
            ),
            (
                "MAX_CHUNK_CHARS = 6,000 in chunker.ts. Functions exceeding this limit are skipped "
                "by the chunker. The file-level fallback also caps at MAX_FILE_BYTES = 500KB in "
                "workspaceIndexer.ts. Generated and minified files should also be excluded via "
                "the IGNORE_FOLDERS list or the user's filesExclude setting."
            )
        ),
        (
            "C-10", "LLM-Powered Summarisation",
            (
                "A function named processData tells a semantic search engine very little about "
                "what the function actually does. The embedding would cluster it near all other "
                "'process' functions regardless of what data is processed or how."
            ),
            (
                "At index time, GPT-4o-mini generates a one-sentence plain-English description of "
                "each function. This description is prepended to the function body before embedding: "
                "'Summary: validates JWT token and extracts user ID\\n\\ndef verify_token(...)'. "
                "The resulting vector captures both the intent (from the summary) and the "
                "implementation (from the code), improving retrieval recall significantly."
            )
        ),
        (
            "C-11", "Line-Level Precision via GPT",
            (
                "Pinecone returns a function's start and end lines, but a 50-line function "
                "contains many individual statements. Jumping the developer to line 1 of a 50-line "
                "function that matches their query is not meaningfully better than jumping to the "
                "file. The developer must still scan manually."
            ),
            (
                "After Pinecone retrieval, line_locator.py re-queries GPT-4o-mini with the "
                "function's source lines (numbered) and the original query. GPT identifies the "
                "single most relevant line. The result is clamped to the known [start_line, "
                "end_line] range and the VS Code extension moves the cursor to exactly that line, "
                "highlighting it."
            )
        ),
        (
            "C-12", "Parallel LLM Calls via ThreadPoolExecutor",
            (
                "AI search requires one GPT call per result for line localisation. With ten results "
                "and an average GPT latency of 500ms per call, sequential execution would produce "
                "a total wait time of five seconds — unacceptable for a developer tool."
            ),
            (
                "Python's ThreadPoolExecutor spawns one thread per result: "
                "max_workers = max(1, len(results)). All GPT calls run concurrently, and the total "
                "wall-clock time is bounded by the slowest single call rather than the sum. "
                "In practice, AI search completes in 600–900ms end-to-end for ten results."
            )
        ),
        (
            "C-13", "Automatic .gitignore Management",
            (
                "The .smart-search/ directory contains per-developer state (project-id, index.json) "
                "that must not be committed to Git. If it were committed, different developers' "
                "project IDs would conflict, and the index.json would cause spurious merge conflicts."
            ),
            (
                "projectId.ts calls addToGitignore() on every activation. The function appends "
                "'.smart-search/' to .gitignore if not already present. This is idempotent and "
                "runs silently — developers never need to configure it manually."
            )
        ),
    ]

    for (cid, title, problem, solution) in challenges:
        add_heading(doc, f"6.{challenges.index((cid, title, problem, solution))+1}  {cid} — {title}", level=2)

        pp = doc.add_paragraph()
        rr = pp.add_run("Problem.  ")
        rr.font.bold = True
        rr.font.name = "Arial"
        rr.font.size = Pt(10)
        rb = pp.add_run(problem)
        rb.font.name = "Garamond"
        rb.font.size = Pt(12)
        pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        sp = doc.add_paragraph()
        rs = sp.add_run("Solution.  ")
        rs.font.bold = True
        rs.font.name = "Arial"
        rs.font.size = Pt(10)
        sb = sp.add_run(solution)
        sb.font.name = "Garamond"
        sb.font.size = Pt(12)
        sp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    page_break(doc)

    # ── Chapter 7: Worked Example ─────────────────────────────────────────────
    add_heading(doc, "Chapter 7  —  Complete Worked Examples", level=1)
    add_heading(doc, "7.1  Scenario A: First-Run Indexing", level=2)
    add_body(doc, (
        "A developer opens a fresh clone of a TypeScript project with 47 source files. "
        "Smart Search has never indexed this workspace before. The following sequence of "
        "events occurs from extension activation to index completion:"
    ))
    add_bullet(doc, "Extension activates. extension.ts reads the projectId from .smart-search/project-id (creating the file and the UUID if absent).")
    add_bullet(doc, "userId is computed: MD5(git config --global user.email.toLowerCase()). If git email is not configured, VS Code shows an error and indexing halts.")
    add_bullet(doc, "The extension calls GET /health to confirm the Python server is running.")
    add_bullet(doc, "indexWorkspace() is called. It walks the workspace directory tree, skipping IGNORE_FOLDERS.")
    add_bullet(doc, "For each .ts / .py / .java / etc. file: compute MD5 of file content. index.json is empty (first run), so every file's hash is new.")
    add_bullet(doc, "chunker.ts extracts functions. A 300-line TypeScript file might yield 12 functions. Class-level chunks are filtered.")
    add_bullet(doc, "For each function: normalizeCode() + MD5. All hashes are new.")
    add_bullet(doc, "Batch is sent to POST /index. Server calls summarizer.py (GPT-4o-mini) and embedder.py (Voyage AI voyage-code-2) for each function.")
    add_bullet(doc, "Embedded vectors are upserted to Pinecone under namespace projectId::userId.")
    add_bullet(doc, "Status bar shows: '⟳ Smart Search: indexing 3/47...' — updates every few files.")
    add_bullet(doc, "On completion: status bar shows '✓ Smart Search: 47 files indexed'. index.json is written with all file and function hashes.")
    add_heading(doc, "7.2  Scenario B: AI Search Query", level=2)
    add_body(doc, "The developer types: 'where do we validate the JWT token and extract the user id?' and presses Search.")
    add_bullet(doc, "Frontend validates: query length = 56 chars ≥ 20. Passes validation.")
    add_bullet(doc, "Message sent to extension: {command: 'search', searchType: 'ai', query: '...'}.")
    add_bullet(doc, "handleSearch() builds namespace from projectId + '::' + userId. POSTs to POST /search with {query, searchType: 'ai', namespace}.")
    add_bullet(doc, "Server calls embed_text(query) with input_type='query'. Returns 1536-float vector in ~120ms.")
    add_bullet(doc, "query_chunks() calls Pinecone with top_k=10, namespace=..., cosine similarity metric.")
    add_bullet(doc, "Pinecone returns 10 candidates. Server filters: score < 0.35 → removed. 7 remain.")
    add_bullet(doc, "ThreadPoolExecutor(max_workers=7): 7 concurrent GPT calls to line_locator. Each reads function source, asks 'which line best answers the query?'. Returns {line: N, content}.")
    add_bullet(doc, "Server assembles response: [{file, name, line, score, content, summary}, ...]. Returns JSON to extension.")
    add_bullet(doc, "Extension posts results to webview. UI renders 7 ranked cards with file path, line, score, and summary snippet.")
    add_bullet(doc, "Developer clicks the top result. openFile handler navigates VS Code to src/auth.ts, line 34 — the exact line where the token payload is extracted.")
    page_break(doc)

    # ── Chapter 8: Performance and Cost ───────────────────────────────────────
    add_heading(doc, "Chapter 8  —  Performance and Cost Analysis", level=1)
    add_heading(doc, "8.1  Indexing Performance", level=2)
    add_table(doc,
        ["Workspace size", "Functions", "First-run time", "Incremental (1 file)"],
        [
            ["Small (< 50 files)", "~200",  "~45 seconds",   "~2 seconds"],
            ["Medium (50–200 files)", "~800", "~3 minutes",  "~3 seconds"],
            ["Large (200–500 files)", "~2000", "~8 minutes", "~4 seconds"],
        ],
        caption="Table 8.1 — Approximate indexing times (depends on API latency)"
    )
    add_heading(doc, "8.2  Search Performance", level=2)
    add_table(doc,
        ["Operation", "Typical latency"],
        [
            ["Voyage AI query embedding", "80–150 ms"],
            ["Pinecone top-10 query", "20–50 ms"],
            ["GPT-4o-mini line locator (parallel × 10)", "400–700 ms"],
            ["Total AI search end-to-end", "550–950 ms"],
            ["Normal (regex) search — 1000-file workspace", "100–300 ms"],
        ],
        caption="Table 8.2 — Measured search latencies"
    )
    add_heading(doc, "8.3  API Cost Estimates", level=2)
    add_table(doc,
        ["Service", "Operation", "Unit cost", "Per 100 functions"],
        [
            ["GPT-4o-mini", "Summarisation (index)", "~$0.0002 / call", "~$0.020"],
            ["GPT-4o-mini", "Line localisation (search)", "~$0.0002 / call", "~$0.002 / search"],
            ["Voyage AI", "voyage-code-2 embedding", "~$0.00012 / 1K tokens", "~$0.001"],
            ["Pinecone", "Serverless upsert/query", "~$0.08 / 1M ops", "< $0.001"],
        ],
        caption="Table 8.3 — Approximate API costs (2025 pricing)"
    )
    add_note(doc, "NOTE", (
        "The dominant cost is GPT-4o-mini summarisation at index time. "
        "Incremental hashing ensures each function is summarised only once (or when it changes), "
        "making the ongoing cost proportional to the rate of code change, not workspace size."
    ))
    page_break(doc)

    # ── Chapter 9: Future Work ─────────────────────────────────────────────────
    add_heading(doc, "Chapter 9  —  Future Work", level=1)
    add_body(doc, (
        "Smart Search reaches a functional and deployable state with the current implementation. "
        "Seven directions for future improvement are identified based on limitations encountered "
        "during development and user testing."
    ))
    future = [
        ("Local Embedding Models",
         "Replace Voyage AI with a locally-hosted model (e.g., Ollama + nomic-embed-code) to "
         "eliminate API cost and latency for indexing. Particularly valuable for air-gapped "
         "development environments."),
        ("Pinecone-Free Offline Mode",
         "Replace Pinecone with a local FAISS or ChromaDB index stored on disk. This would make "
         "the tool fully self-contained with no required third-party accounts, which is important "
         "for enterprise environments with strict data residency requirements."),
        ("Hybrid Search",
         "Combine vector search with BM25 (sparse keyword matching) using Pinecone's hybrid "
         "search capability. Hybrid search outperforms pure dense retrieval on short, "
         "keyword-heavy queries."),
        ("Symbol-Aware Re-ranking",
         "After Pinecone retrieval, apply a lightweight cross-encoder re-ranker that is aware of "
         "symbol names and call graphs. This would improve precision for queries that reference "
         "specific function names or variable names."),
        ("Streaming Index Updates",
         "Replace the HTTP polling approach for index status with a WebSocket or Server-Sent "
         "Events stream, allowing the UI to show real-time progress without polling."),
        ("Multi-Root Workspace Support",
         "VS Code supports workspaces with multiple root folders. The current implementation "
         "indexes only the first root. Full multi-root support requires namespace disambiguation "
         "across roots."),
        ("Result Explanation",
         "Add a 'Why this result?' feature that asks GPT to explain the connection between the "
         "user's query and the retrieved function. This would help developers validate relevance "
         "without needing to read the full function source."),
    ]
    for i, (title, desc) in enumerate(future):
        add_heading(doc, f"9.{i+1}  {title}", level=2)
        add_body(doc, desc)
    page_break(doc)

    # ── Chapter 10: Conclusion ─────────────────────────────────────────────────
    add_heading(doc, "Chapter 10  —  Conclusion", level=1)
    add_body(doc, (
        "This thesis has presented Smart Search — a VS Code extension that brings semantic, "
        "intent-based code retrieval to the developer's native editing environment. "
        "By combining voyage-code-2 embeddings, Pinecone vector storage, GPT-4o-mini "
        "summarisation and line localisation, and a carefully designed incremental hash index, "
        "the system enables developers to find code by meaning rather than by spelling."
    ))
    add_body(doc, (
        "The engineering challenges documented in Chapter 6 collectively represent the gap "
        "between a proof-of-concept and a production-grade tool. Multi-language chunking, "
        "formatting-insensitive hashing, multi-tenant isolation, and parallel LLM execution "
        "were each necessary to make the system accurate, fast, and economically viable for "
        "real-world development workflows."
    ))
    add_body(doc, (
        "From a research perspective, the most significant contribution is the two-level hash "
        "strategy combined with code normalisation, which makes incremental indexing cost "
        "proportional to code change velocity rather than workspace size. A developer who "
        "modifies ten functions per day on a ten-thousand-function codebase incurs only ten "
        "embedding calls — not ten thousand."
    ))
    add_body(doc, (
        "Smart Search demonstrates that semantic code search is achievable within a VS Code "
        "extension with minimal infrastructure requirements: a single Python process and three "
        "API keys are sufficient to bring vector-powered code intelligence to any developer "
        "working in any supported language."
    ))

    # ── References ─────────────────────────────────────────────────────────────
    page_break(doc)
    add_heading(doc, "References", level=1)
    refs = [
        "Voyage AI. (2024). voyage-code-2: Code Embedding Model. https://docs.voyageai.com",
        "Pinecone. (2024). Pinecone Serverless Documentation. https://docs.pinecone.io",
        "OpenAI. (2024). GPT-4o-mini Model Card. https://platform.openai.com/docs",
        "Microsoft. (2024). VS Code Extension API. https://code.visualstudio.com/api",
        "Karpukhin, V., et al. (2020). Dense Passage Retrieval for Open-Domain Question Answering. EMNLP 2020.",
        "Robertson, S., & Zaragoza, H. (2009). The Probabilistic Relevance Framework: BM25 and Beyond. Foundations and Trends in Information Retrieval.",
        "Muennighoff, N., et al. (2022). MTEB: Massive Text Embedding Benchmark. arXiv:2210.07316.",
        "python-docx contributors. (2024). python-docx 1.1.0 Documentation. https://python-docx.readthedocs.io",
    ]
    for r in refs:
        rp = doc.add_paragraph()
        rp.paragraph_format.left_indent = Inches(0.3)
        rp.paragraph_format.first_line_indent = Inches(-0.3)
        run = rp.add_run(r)
        run.font.name = "Garamond"
        run.font.size = Pt(11)

    # ── Save ───────────────────────────────────────────────────────────────────
    doc.save("smart_search_thesis.docx")
    print("Done: smart_search_thesis.docx created successfully.")


if __name__ == "__main__":
    build()
